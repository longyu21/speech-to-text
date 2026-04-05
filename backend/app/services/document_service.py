from io import BytesIO
from typing import Sequence

from docx import Document


def build_transcript_docx(title: str, sections: Sequence[tuple[str, str]]) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    for index, (section_title, section_text) in enumerate(sections):
        if index > 0:
            document.add_paragraph("")
        document.add_heading(section_title, level=2)
        document.add_paragraph(section_text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()
